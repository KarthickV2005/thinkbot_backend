import nltk
nltk.download("punkt")
nltk.download("stopwords")
nltk.download("wordnet")


import re
import docx
from docx import Document
from odf import text, teletype
from odf.opendocument import load

from nltk.corpus import stopwords
from nltk.stem import PorterStemmer, WordNetLemmatizer

class ProcessorAgent:
    def __init__(self):
        self.stop_words = set(stopwords.words("english"))
        self.stemmer = PorterStemmer()
        self.lemmatizer = WordNetLemmatizer()

    def preprocess_text(self, text_data: str) -> str:
        """
        Preprocess raw text:
        1) Lowercase
        2) Remove special characters
        3) Tokenization
        4) Remove stopwords
        5) Stemming
        6) Lemmatization
        """
        # Lowercase
        text_data = text_data.lower()

        # Remove special characters (keep alphanumeric + spaces)
        text_data = re.sub(r"[^a-z0-9\s]", "", text_data)

        # Tokenize
        tokens = nltk.word_tokenize(text_data)

        # Remove stopwords
        tokens = [word for word in tokens if word not in self.stop_words]

        # Apply stemming
        stemmed = [self.stemmer.stem(word) for word in tokens]

        # Apply lemmatization
        lemmatized = [self.lemmatizer.lemmatize(word) for word in stemmed]

        # Join back into string
        return " ".join(lemmatized)

    def extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def extract_text_from_odt(self, file_path: str) -> str:
        odt_doc = load(file_path)
        all_text = ""
        for elem in odt_doc.getElementsByType(text.P):
            all_text += teletype.extractText(elem) + "\n"
        return all_text

    def overwrite_file(self, file_path: str, content: str):
        """
        Overwrite the original file with preprocessed text
        """
        if file_path.endswith(".txt"):
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

        elif file_path.endswith(".docx"):
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)

        elif file_path.endswith(".odt"):
            from odf.opendocument import OpenDocumentText
            from odf.text import P
            odt_doc = OpenDocumentText()
            odt_doc.text.addElement(P(text=content))
            odt_doc.save(file_path)

        else:
            raise ValueError("Unsupported file format for overwrite")

    def process_file(self, file_path: str) -> str:
        """
        Main file handler: preprocess + overwrite the same file
        """
        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        elif file_path.endswith(".docx"):
            raw_text = self.extract_text_from_docx(file_path)
        elif file_path.endswith(".odt"):
            raw_text = self.extract_text_from_odt(file_path)
        else:
            raise ValueError("Unsupported file format")

        print("\n===== BEFORE (Original Content) =====\n")
        print(raw_text)

        preprocessed = self.preprocess_text(raw_text)

        print("\n===== AFTER (Preprocessed Content) =====\n")
        print(preprocessed)


        
        self.overwrite_file(file_path, preprocessed)

        print(f"✅ File '{file_path}' has been overwritten with preprocessed text.")
        return preprocessed


if __name__ == "__main__":
    agent = ProcessorAgent()
    agent.process_file(r"S:\ThinkBot-new\backend\text_files\ecommerce.txt")